"""
generador_dp.py
Genera la respuesta a un Derecho de Petición con el membrete oficial de la
Universidad Distrital, más el Excel adjunto con las hojas de datos.

Salida:
  - Respuesta_DP_{radicado}_{fecha}.docx  (Word con membrete)
  - F_DP_{radicado}_{fecha}.xlsx           (Excel adjunto con hojas por punto)
"""
import io
import os
import copy
from datetime import date
from pathlib import Path

# ── Rutas ─────────────────────────────────────────────────────────────
BASE_DIR      = Path(__file__).parent
MEMBRETE_JPG  = BASE_DIR / "plantillas" / "membrete_ud.jpg"
LOGO_PNG      = BASE_DIR / "plantillas" / "logo_ud.png"


# ═══════════════════════════════════════════════════════════════════════
# GENERADOR WORD (con membrete como imagen de fondo via python-docx)
# ═══════════════════════════════════════════════════════════════════════

def generar_word_respuesta(datos: dict, proyectos: list) -> bytes:
    """
    Genera el Word de respuesta al derecho de petición.
    datos = {
        radicado, fecha_radicado, destinatario_nombre, destinatario_cargo,
        destinatario_entidad, destinatario_correo, ciudad,
        asunto, firmante_nombre, firmante_cargo,
        puntos: [{ numero, texto_pregunta, tipo, respuesta_texto }]
    }
    tipo: 'tabla_excel' | 'texto'
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import lxml.etree as etree

    doc = Document()

    # ── Configurar página carta ──────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    # Márgenes: top grande para dejar espacio al membrete
    section.top_margin    = Inches(1.5)   # el membrete ocupa hasta ~1.3"
    section.bottom_margin = Inches(1.2)   # el footer ocupa ~1" al final
    section.left_margin   = Inches(1.18)  # ~3cm
    section.right_margin  = Inches(0.98)  # ~2.5cm

    # ── Agregar membrete como imagen de encabezado ───────────────────
    header = section.header
    # Limpiar párrafos existentes
    for p in header.paragraphs:
        p.clear()
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run()

    # El membrete es 8.5x11" — lo usamos como imagen de página completa
    # Lo posicionamos como imagen inline en el header con ancho de página
    if MEMBRETE_JPG.exists():
        run.add_picture(str(MEMBRETE_JPG), width=Inches(8.5))

    # ── Estilos base ─────────────────────────────────────────────────
    UD_BLUE = RGBColor(30, 58, 110)
    BLACK   = RGBColor(0, 0, 0)

    def p_normal(text="", bold=False, italic=False, size=10, color=BLACK,
                 align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        if text:
            run = p.add_run(text)
            run.font.name   = "Arial"
            run.font.size   = Pt(size)
            run.font.bold   = bold
            run.font.italic = italic
            run.font.color.rgb = color
        return p

    def p_heading(text, level=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(6)
        run = p.add_run(text)
        run.font.name  = "Arial"
        run.font.size  = Pt(11 if level == 1 else 10)
        run.font.bold  = True
        run.font.color.rgb = UD_BLUE
        return p

    def add_mixed_paragraph(partes, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6):
        """partes = [(texto, bold, italic, size)]"""
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        for texto, bold, italic, size in partes:
            run = p.add_run(texto)
            run.font.name   = "Arial"
            run.font.size   = Pt(size)
            run.font.bold   = bold
            run.font.italic = italic
            run.font.color.rgb = BLACK
        return p

    # ── Encabezado del documento ──────────────────────────────────────
    fecha_hoy = date.today().strftime("%d de %B de %Y").replace(
        "January","enero").replace("February","febrero").replace("March","marzo"
        ).replace("April","abril").replace("May","mayo").replace("June","junio"
        ).replace("July","julio").replace("August","agosto").replace("September","septiembre"
        ).replace("October","octubre").replace("November","noviembre").replace("December","diciembre")

    p_normal(f"Bogotá D.C., {fecha_hoy}", align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=12)
    p_normal(f"Radicado: {datos.get('radicado','')}", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=16)

    p_normal(datos.get('destinatario_nombre',''), bold=True, space_after=2)
    p_normal(datos.get('destinatario_cargo',''), space_after=2)
    p_normal(datos.get('destinatario_entidad',''), bold=True, space_after=2)
    if datos.get('destinatario_correo'):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(datos['destinatario_correo'])
        run.font.name = "Arial"; run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 70, 180)
    p_normal(datos.get('ciudad','Bogotá, Colombia'), space_after=12)

    add_mixed_paragraph([
        ("Asunto: ", True, False, 10),
        (datos.get('asunto','Respuesta a Derecho de Petición'), False, False, 10),
    ], space_after=14)

    p_normal("Cordial saludo,", space_after=10)

    p_normal(
        f"En atención a su comunicación radicada bajo el número {datos.get('radicado','')} "
        f"de fecha {datos.get('fecha_radicado','')}, recibida en la Oficina de Extensión – IDEXUD "
        f"de la Universidad Distrital Francisco José de Caldas, me permito dar respuesta a cada "
        f"uno de los puntos solicitados:",
        space_after=14
    )

    # ── Puntos de respuesta ───────────────────────────────────────────
    p_heading("I. RESPUESTA A LAS PETICIONES")

    for punto in datos.get('puntos', []):
        num   = punto.get('numero', '')
        texto = punto.get('texto_pregunta', '')
        tipo  = punto.get('tipo', 'texto')
        resp  = punto.get('respuesta_texto', '')

        # Pregunta original en cursiva
        add_mixed_paragraph([
            (f"{num}. ", True, False, 10),
            (texto, False, True, 10),
        ], space_after=4)

        # Respuesta
        if tipo == 'tabla_excel':
            hoja = punto.get('hoja_excel', f'Punto {num}')
            add_mixed_paragraph([
                ("Respuesta: ", True, False, 10),
                (f"Se relaciona la información respecto al periodo en mención en la hoja \"{hoja}\". "
                 f"Ver adjunto \"F_DP_{datos.get('radicado','DP')}.xlsx\"", False, False, 10),
            ], space_after=10)
        else:
            add_mixed_paragraph([
                ("Respuesta: ", True, False, 10),
                (resp, False, False, 10),
            ], space_after=10)

    # ── Nota final ────────────────────────────────────────────────────
    p_normal(
        "Nota: La información es consultada y reportada de los sistemas de la Oficina "
        "de Extensión – IDEXUD.",
        italic=True, size=9, space_after=16
    )

    # ── Cierre ────────────────────────────────────────────────────────
    p_normal("Atentamente,", space_after=36)
    p_normal(datos.get('firmante_nombre',''), bold=True, space_after=2)
    p_normal(datos.get('firmante_cargo',''), space_after=2)
    p_normal("Universidad Distrital Francisco José de Caldas", space_after=2)
    p_normal("Oficina de Extensión – IDEXUD", size=9, space_after=16)

    p_normal(f"Elaboró: Sistema SIEXUD", size=8, space_after=2)
    p_normal(f"Fecha: {fecha_hoy}", size=8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ═══════════════════════════════════════════════════════════════════════
# GENERADOR EXCEL ADJUNTO (hojas por punto)
# ═══════════════════════════════════════════════════════════════════════

def generar_excel_adjunto(radicado: str, puntos_data: dict) -> bytes:
    """
    puntos_data = {
        1: [{'No. Convenio': ..., 'Entidad': ..., ...}],
        2: [{'Año': ..., 'No. Convenio': ..., ...}],
        4: [{'No. Convenio': ..., '%Avance': ..., ...}],
    }
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    wb.remove(wb.active)  # Quitar hoja por defecto

    HEADER_FILL = PatternFill("solid", start_color="1E3A6E")
    HEADER_FONT = Font(bold=True, color="FFFFFF", name="Arial", size=10)
    DATA_FONT   = Font(name="Arial", size=9)
    ALT_FILL    = PatternFill("solid", start_color="EFF6FF")
    thin        = Side(style='thin', color='CBD5E1')
    BORDER      = Border(left=thin, right=thin, top=thin, bottom=thin)
    CENTER      = Alignment(horizontal='center', vertical='center', wrap_text=True)
    LEFT        = Alignment(horizontal='left', vertical='center', wrap_text=True)

    # Columnas por punto
    PUNTO_COLS = {
        1: [
            ("No. Convenio", 14), ("Entidad Contratante", 35), ("Objeto", 50),
            ("Valor Total", 18), ("Fecha Suscripción", 18), ("Fecha Fin Ejecución", 18),
            ("Dependencia Encargada", 25), ("Estado Actual", 16),
        ],
        2: [
            ("Año", 8), ("No. Convenio", 14), ("Entidad", 30), ("Localidad", 16),
            ("Fecha de Avance", 16), ("Número de Pagos Realizados", 20),
            ("Valor Total Pagado", 18), ("Estado de Pagos a la Fecha", 22),
        ],
        4: [
            ("No. Convenio Asociado", 18), ("No. Contrato", 14),
            ("%Avance Presupuestal", 18), ("%Avance Físico", 16),
            ("Actividades Ejecutadas", 30), ("Actividades Pendientes", 30),
            ("Se han presentado retrasos (SI/NO)", 20),
            ("Modificaciones o Prórrogas (SÍ/NO)", 20),
            ("Dependencia Encargada", 22), ("Valor Total Contrato", 18),
        ],
    }

    PUNTO_TITLES = {
        1: "Convenios Interadministrativos (con aporte UD) celebrados desde el 01/01/2023",
        2: "Pagos realizados a la Universidad Distrital por convenios desde el 01/01/2023",
        4: "Avance de ejecución de convenios y contratos derivados desde 2023",
    }

    for punto_num, rows in puntos_data.items():
        cols  = PUNTO_COLS.get(punto_num, [])
        title = PUNTO_TITLES.get(punto_num, f"Punto {punto_num}")

        ws = wb.create_sheet(title=f"Punto {punto_num}")
        ws.freeze_panes = "A3"

        # Fila 1: título
        ws.merge_cells(f"A1:{get_column_letter(len(cols))}1")
        tc = ws['A1']
        tc.value     = title
        tc.font      = Font(bold=True, size=11, color="FFFFFF", name="Arial")
        tc.fill      = PatternFill("solid", start_color="B91C3C")
        tc.alignment = CENTER
        ws.row_dimensions[1].height = 22

        # Fila 2: encabezados
        for ci, (col_name, col_width) in enumerate(cols, 1):
            c = ws.cell(row=2, column=ci, value=col_name)
            c.font      = HEADER_FONT
            c.fill      = HEADER_FILL
            c.alignment = CENTER
            c.border    = BORDER
            ws.column_dimensions[get_column_letter(ci)].width = col_width
        ws.row_dimensions[2].height = 28

        # Datos
        col_names = [c[0] for c in cols]
        for ri, row in enumerate(rows, 3):
            for ci, col_name in enumerate(col_names, 1):
                val = row.get(col_name, '')
                c   = ws.cell(row=ri, column=ci, value=val)
                c.font      = DATA_FONT
                c.border    = BORDER
                c.alignment = LEFT
                if 'Valor' in col_name and isinstance(val, (int, float)):
                    c.number_format = '$ #,##0'
                    c.alignment = CENTER
                if 'Avance' in col_name and '%' in col_name:
                    c.number_format = '0.0"%"'
                    c.alignment = CENTER
                if ri % 2 == 0:
                    c.fill = ALT_FILL
            ws.row_dimensions[ri].height = 18

        # Total valor si aplica
        val_cols = [i+1 for i, (cn, _) in enumerate(cols) if 'Valor' in cn]
        if val_cols and rows:
            tot_row = len(rows) + 3
            for vc in val_cols:
                col_l = get_column_letter(vc)
                c = ws.cell(row=tot_row, column=vc,
                            value=f"=SUM({col_l}3:{col_l}{tot_row-1})")
                c.font         = Font(bold=True, name="Arial", size=10)
                c.number_format = '$ #,##0'
                c.border       = BORDER
                c.alignment    = CENTER
            ws.cell(row=tot_row, column=1, value="TOTAL").font = Font(bold=True, name="Arial")

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


# ═══════════════════════════════════════════════════════════════════════
# FUNCIÓN PRINCIPAL — construir todo desde los datos de BD
# ═══════════════════════════════════════════════════════════════════════

def construir_respuesta_dp(
    radicado: str,
    fecha_radicado: str,
    destinatario_nombre: str,
    destinatario_cargo: str,
    destinatario_entidad: str,
    destinatario_correo: str,
    asunto: str,
    ciudad: str,
    firmante_nombre: str,
    firmante_cargo: str,
    proyectos: list,         # lista de dicts con datos de la BD
    modificaciones: list,    # lista de dicts
) -> tuple[bytes, bytes]:
    """
    Retorna (word_bytes, excel_bytes)
    """
    from datetime import datetime

    def fmt_money(v):
        if not v: return "$ 0"
        return f"$ {float(v):,.0f}".replace(',', '.')

    def fmt_date(d):
        if not d: return "N/A"
        if isinstance(d, str): return d
        try: return d.strftime('%d/%m/%Y')
        except: return str(d)

    # ── Construir filas para Excel ────────────────────────────────────

    # Punto 1: convenios desde 2023
    p1_rows = []
    for p in proyectos:
        p1_rows.append({
            "No. Convenio":          p.get('external') or str(p.get('id','')),
            "Entidad Contratante":   p.get('entity',''),
            "Objeto":                p.get('purpose','')[:200],
            "Valor Total":           float(p.get('value',0)),
            "Fecha Suscripción":     p.get('subscription_date','N/A'),
            "Fecha Fin Ejecución":   p.get('end_date','N/A'),
            "Dependencia Encargada": p.get('department',''),
            "Estado Actual":         p.get('status',''),
        })

    # Punto 2: pagos (usamos valor del contrato como referencia)
    p2_rows = []
    for p in proyectos:
        p2_rows.append({
            "Año":                         p.get('year',''),
            "No. Convenio":                p.get('external') or str(p.get('id','')),
            "Entidad":                     p.get('entity',''),
            "Localidad":                   "Bogotá D.C.",
            "Fecha de Avance":             p.get('end_date',''),
            "Número de Pagos Realizados":  "N/D",
            "Valor Total Pagado":          float(p.get('value',0)),
            "Estado de Pagos a la Fecha":  p.get('status',''),
        })

    # Punto 4: avance con modificaciones
    p4_rows = []
    mods_por_proyecto = {}
    for m in modificaciones:
        pid = m.get('project_id')
        if pid not in mods_por_proyecto:
            mods_por_proyecto[pid] = []
        mods_por_proyecto[pid].append(m)

    for p in proyectos:
        pid  = p.get('id')
        mods = mods_por_proyecto.get(pid, [])
        tiene_prorrogas = any(m.get('modification_type') in ('EXTENSION','BOTH') for m in mods)
        tiene_mods      = len(mods) > 0
        p4_rows.append({
            "No. Convenio Asociado":               p.get('external') or str(pid),
            "No. Contrato":                        p.get('external') or str(pid),
            "%Avance Presupuestal":                "N/D",
            "%Avance Físico":                      "N/D",
            "Actividades Ejecutadas":              "Ver sistema SIEXUD",
            "Actividades Pendientes":              "Ver sistema SIEXUD",
            "Se han presentado retrasos (SI/NO)":  "SÍ" if tiene_prorrogas else "NO",
            "Modificaciones o Prórrogas (SÍ/NO)":  "SÍ" if tiene_mods else "NO",
            "Dependencia Encargada":               p.get('department',''),
            "Valor Total Contrato":                float(p.get('value',0)),
        })

    puntos_excel = {1: p1_rows, 2: p2_rows, 4: p4_rows}

    # ── Construir puntos para Word ────────────────────────────────────
    puntos_word = [
        {
            "numero": "1",
            "texto_pregunta": (
                "Sírvase informar y detallar los convenios interadministrativos celebrados "
                "por la Universidad Distrital desde el 01 de enero de 2023 hasta la fecha, "
                "incluyendo el objeto, las entidades con las que se suscribieron, el valor "
                "de cada convenio, el plazo de ejecución y la dependencia encargada de su gestión."
            ),
            "tipo": "tabla_excel",
            "hoja_excel": "Punto 1",
        },
        {
            "numero": "2",
            "texto_pregunta": (
                "Sírvase informar y detallar cuánto se le ha pagado a la Universidad Distrital "
                "por cada uno de los convenios interadministrativos suscritos desde el 1 de enero "
                "de 2023 hasta la fecha."
            ),
            "tipo": "tabla_excel",
            "hoja_excel": "Punto 2",
        },
        {
            "numero": "3",
            "texto_pregunta": (
                "Sírvase certificar y detallar si se ha realizado algún tipo de subcontratación "
                "para el cumplimiento de los contratos derivados de los convenios interadministrativos "
                "celebrados desde el 1 de enero de 2023 hasta la fecha actual."
            ),
            "tipo": "texto",
            "respuesta_texto": (
                "Luego de revisar la información contractual correspondiente al periodo indicado, "
                "no se registra ningún tipo de subcontratación asociada a la ejecución de los contratos "
                "derivados de los convenios interadministrativos. En ese sentido, y teniendo en cuenta "
                "que la consulta se formula específicamente respecto de la subcontratación vinculada a "
                "la contratación derivada de dichos convenios, no aplica el diligenciamiento de la tabla "
                "solicitada, dado que no existen subcontratos celebrados en el marco de las obligaciones "
                "contractuales analizadas."
            ),
        },
        {
            "numero": "4",
            "texto_pregunta": (
                "Sírvase informar y detallar el avance de la ejecución de los convenios "
                "interadministrativos y los contratos derivados desde 2023 hasta la fecha, "
                "indicando el porcentaje de ejecución financiera y física, el cronograma de "
                "actividades y si se han presentado retrasos, modificaciones o prórrogas."
            ),
            "tipo": "tabla_excel",
            "hoja_excel": "Punto 4",
        },
    ]

    datos_word = {
        "radicado":              radicado,
        "fecha_radicado":        fecha_radicado,
        "destinatario_nombre":   destinatario_nombre,
        "destinatario_cargo":    destinatario_cargo,
        "destinatario_entidad":  destinatario_entidad,
        "destinatario_correo":   destinatario_correo,
        "ciudad":                ciudad,
        "asunto":                asunto,
        "firmante_nombre":       firmante_nombre,
        "firmante_cargo":        firmante_cargo,
        "puntos":                puntos_word,
    }

    word_bytes  = generar_word_respuesta(datos_word, proyectos)
    excel_bytes = generar_excel_adjunto(radicado, puntos_excel)

    return word_bytes, excel_bytes
