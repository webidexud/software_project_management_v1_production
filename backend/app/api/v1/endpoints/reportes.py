# backend/app/api/v1/endpoints/reportes.py — v2.0
# MEJORAS:
#  - Detección de intención: analítico vs generación de archivo
#  - Genera Excel, PDF o Word según lo que pida el usuario
#  - Pasa datos individuales de proyectos (no solo totales)
#  - Optimización de tokens: contexto dinámico según la pregunta
#  - Endpoint /reportes/download/{token} para descargar archivos generados
import io
import os
import re
import json
import uuid
import httpx
import tempfile
from datetime import date, datetime
from typing import Optional
from pathlib import Path

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse, Response
from sqlalchemy.orm import Session
from sqlalchemy import func
from pydantic import BaseModel

from app.db.database import get_db
from app.core.config import settings
from app.models.catalogs import (
    Project, ProjectStatus, Entity, ExecutingDepartment,
    ExecutionModality, FinancingType
)
from app.models.project_modification import ProjectModification
from app.models.project_document import ProjectDocument

router = APIRouter(prefix="/reportes", tags=["Reportes"])

# ── Almacén temporal de archivos generados (en memoria) ───────────────
# { token: { "filename": str, "content": bytes, "mimetype": str } }
_file_store: dict = {}   # { token: { content, mimetype, filename, created_at } }
_FILE_TTL_SECONDS = 1800  # 30 minutos


# ── Helpers ───────────────────────────────────────────────────────────

def apply_filters(q, year=None, status_ids=None, valor_min=None, valor_max=None,
                  fecha_inicio_desde=None, fecha_inicio_hasta=None,
                  fecha_fin_desde=None, fecha_fin_hasta=None,
                  entidad_id=None, departamento_id=None):
    from datetime import date as date_type
    if year:       q = q.filter(Project.project_year == int(year))
    # Múltiples estados (lista separada por comas)
    if status_ids:
        ids = [int(x) for x in str(status_ids).split(',') if x.strip()]
        if ids:
            q = q.filter(Project.project_status_id.in_(ids))
    if valor_min:  q = q.filter(Project.project_value >= float(valor_min))
    if valor_max:  q = q.filter(Project.project_value <= float(valor_max))
    # Rango fecha inicio
    if fecha_inicio_desde:
        q = q.filter(Project.start_date >= fecha_inicio_desde)
    if fecha_inicio_hasta:
        q = q.filter(Project.start_date <= fecha_inicio_hasta)
    # Rango fecha fin
    if fecha_fin_desde:
        q = q.filter(Project.end_date >= fecha_fin_desde)
    if fecha_fin_hasta:
        q = q.filter(Project.end_date <= fecha_fin_hasta)
    if entidad_id:
        q = q.filter(Project.entity_id == int(entidad_id))
    if departamento_id:
        q = q.filter(Project.executing_department_id == int(departamento_id))
    return q


def get_status_name(db, status_id):
    st = db.query(ProjectStatus).filter(ProjectStatus.status_id == status_id).first()
    return st.status_name if st else '—'


def get_entity_name(db, entity_id):
    en = db.query(Entity).filter(Entity.entity_id == entity_id).first()
    return en.entity_name if en else '—'


def get_dept_name(db, dept_id):
    dp = db.query(ExecutingDepartment).filter(ExecutingDepartment.department_id == dept_id).first()
    return dp.department_name if dp else '—'


def fmt_money(v):
    if v is None: return '—'
    return f"${float(v):,.0f}".replace(',', '.')


def fmt_date(d):
    if not d: return '—'
    if isinstance(d, str): return d
    return d.strftime('%d/%m/%Y')


# ── Detección de intención ────────────────────────────────────────────

INTENT_PATTERNS = {
    'excel': [
        r'excel', r'hoja de c[aá]lculo', r'xlsx', r'spreadsheet',
        r'tabla en excel', r'exportar', r'descargar.*tabla',
    ],
    'pdf': [
        r'pdf', r'reporte.*pdf', r'informe.*pdf', r'documento.*pdf',
        r'generar.*reporte', r'crear.*reporte', r'dame.*reporte',
        r'necesito.*reporte', r'reporte.*con.*tabla', r'informe',
    ],
    'word': [
        r'word', r'docx', r'documento word', r'\.doc\b',
    ],
}

REPORT_KEYWORDS = [
    r'reporte', r'informe', r'lista de', r'dame.*proyectos',
    r'mu[eé]strame.*proyectos', r'todos los proyectos', r'proyectos.*de \d{4}',
    r'generar', r'crear.*lista', r'exportar', r'descargar',
    r'tabla.*con', r'necesito.*datos',
]

FILTER_PATTERNS = {
    'year':   r'\b(20\d{2})\b',
    'status': r'(en ejecuci[oó]n|activo|terminado|liquidado|suspendido|cerrado|en proceso)',
}

STATUS_MAP = {
    'en ejecucion': 'EN EJECUCIÓN',
    'en ejecución': 'EN EJECUCIÓN',
    'activo': 'ACTIVO',
    'terminado': 'TERMINADO',
    'liquidado': 'LIQUIDADO',
    'suspendido': 'SUSPENDIDO',
    'cerrado': 'CERRADO',
    'en proceso': 'EN PROCESO',
}


def detect_intent(message: str) -> dict:
    """Detecta si el usuario quiere un archivo y de qué tipo, con qué filtros."""
    msg_lower = message.lower()

    # Tipo de archivo
    file_type = None
    for ftype, patterns in INTENT_PATTERNS.items():
        for p in patterns:
            if re.search(p, msg_lower):
                file_type = ftype
                break
        if file_type:
            break

    # Si menciona keywords de reporte sin tipo específico → PDF por defecto
    is_report = any(re.search(p, msg_lower) for p in REPORT_KEYWORDS)
    if is_report and not file_type:
        file_type = 'pdf'

    # Filtros detectados
    filters = {}
    year_match = re.search(FILTER_PATTERNS['year'], msg_lower)
    if year_match:
        filters['year'] = year_match.group(1)

    status_match = re.search(FILTER_PATTERNS['status'], msg_lower)
    if status_match:
        filters['status_name'] = STATUS_MAP.get(status_match.group(1).lower(), status_match.group(1))

    # Columnas pedidas explícitamente
    col_keywords = {
        'año': 'year', 'numero': 'number', 'número': 'number',
        'objeto': 'purpose', 'nombre': 'name', 'entidad': 'entity',
        'valor': 'value', 'estado': 'status', 'fecha': 'dates',
        'dependencia': 'department', 'modificaciones': 'mods',
    }
    requested_cols = []
    for kw, col in col_keywords.items():
        if kw in msg_lower:
            requested_cols.append(col)

    return {
        'file_type': file_type,
        'is_report':  is_report or bool(file_type),
        'filters':    filters,
        'columns':    requested_cols if requested_cols else None,
    }


# ── Obtener proyectos con filtros de lenguaje natural ────────────────

def get_projects_for_report(db, intent_filters: dict, extra_params: dict = None):
    """Consulta proyectos aplicando filtros del intent y los filtros del panel."""
    q = db.query(Project).filter(Project.is_active == True)

    # Filtros del panel de reportes
    if extra_params:
        if extra_params.get('year'):      q = q.filter(Project.project_year == int(extra_params['year']))
        if extra_params.get('valor_min'): q = q.filter(Project.project_value >= float(extra_params['valor_min']))
        if extra_params.get('valor_max'): q = q.filter(Project.project_value <= float(extra_params['valor_max']))
        if extra_params.get('status_id'): q = q.filter(Project.project_status_id == int(extra_params['status_id']))

    # Filtros del lenguaje natural
    if intent_filters.get('year'):
        q = q.filter(Project.project_year == int(intent_filters['year']))

    if intent_filters.get('status_name'):
        status = db.query(ProjectStatus).filter(
            func.upper(ProjectStatus.status_name).contains(
                intent_filters['status_name'].upper()
            )
        ).first()
        if status:
            q = q.filter(Project.project_status_id == status.status_id)

    return q.order_by(Project.project_year.desc(), Project.project_id).all()


# ── Preparar datos de proyectos (con caché de nombres) ───────────────

def build_project_rows(db, proyectos, columns=None):
    """Construye lista de dicts con los datos de cada proyecto."""
    # Caché para evitar N+1 queries
    status_cache = {}
    entity_cache = {}
    dept_cache   = {}

    rows = []
    for p in proyectos:
        if p.project_status_id not in status_cache:
            status_cache[p.project_status_id] = get_status_name(db, p.project_status_id)
        if p.entity_id not in entity_cache:
            entity_cache[p.entity_id] = get_entity_name(db, p.entity_id)
        if p.executing_department_id not in dept_cache:
            dept_cache[p.executing_department_id] = get_dept_name(db, p.executing_department_id)

        row = {
            'id':         p.project_id,
            'year':       p.project_year,
            'external':   p.external_project_number or '—',
            'name':       p.project_name,
            'purpose':    p.project_purpose,
            'status':     status_cache[p.project_status_id],
            'entity':     entity_cache[p.entity_id],
            'department': dept_cache[p.executing_department_id],
            'value':      float(p.project_value or 0),
            'start_date': fmt_date(p.start_date),
            'end_date':   fmt_date(p.end_date),
            'benefit':    float(p.institutional_benefit_value or 0),
            'email':      p.main_email or '—',
            'act':        p.administrative_act or '—',
        }
        rows.append(row)
    return rows


# ── Generadores de archivos ───────────────────────────────────────────

def generate_excel(rows, title, columns=None):
    """
    Retorna un Workbook (no bytes).
    Llamar .save(buf) para serializar.
    Cuando viene del chat (columns != None) genera solo 1 hoja simple.
    Cuando viene del botón global genera 4 hojas completas.
    """
    if columns:
        # Versión simple de 1 hoja para el chat IA
        return _generate_excel_simple(rows, title, columns)
    return _generate_excel_multisheet(rows, title)


def _generate_excel_simple(rows, title, columns=None):
    """Excel de 1 hoja — para reportes generados desde el chat."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    ALL_COLS = [
        ('id',         '#ID',                7),
        ('year',       'Año',                6),
        ('external',   'N° Externo',        16),
        ('name',       'Nombre',            48),
        ('purpose',    'Objeto',            48),
        ('status',     'Estado',            18),
        ('entity',     'Entidad',           32),
        ('department', 'Dependencia',       26),
        ('value',      'Valor Contrato',    18),
        ('start_date', 'Fecha Inicio',      14),
        ('end_date',   'Fecha Fin',         14),
    ]

    col_map = {
        'year':'year','number':'external','purpose':'purpose','name':'name',
        'entity':'entity','value':'value','status':'status',
        'dates':'start_date','department':'department',
    }
    if columns:
        keys = {col_map[c] for c in columns if c in col_map}
        keys.add('id')
        use_cols = [c for c in ALL_COLS if c[0] in keys]
    else:
        use_cols = ALL_COLS

    wb  = Workbook()
    ws  = wb.active
    ws.title = "Proyectos"
    ws.freeze_panes = "A3"

    H_FILL = PatternFill("solid", start_color="1E3A6E")
    H_FONT = Font(bold=True, color="FFFFFF", name="Arial", size=10)
    D_FONT = Font(name="Arial", size=9)
    ALT    = PatternFill("solid", start_color="EFF6FF")
    thin   = Side(style="thin", color="D1D5DB")
    BDR    = Border(left=thin, right=thin, top=thin, bottom=thin)
    CTR    = Alignment(horizontal="center", vertical="center")
    LEFT   = Alignment(horizontal="left", vertical="center", wrap_text=True)

    # Título
    ws.merge_cells(f"A1:{get_column_letter(len(use_cols))}1")
    t = ws["A1"]
    t.value = title
    t.font  = Font(bold=True, size=12, color="FFFFFF", name="Arial")
    t.fill  = PatternFill("solid", start_color="B91C3C")
    t.alignment = CTR
    ws.row_dimensions[1].height = 24

    # Encabezados
    for ci, (key, label, width) in enumerate(use_cols, 1):
        c = ws.cell(row=2, column=ci, value=label)
        c.font = H_FONT; c.fill = H_FILL; c.alignment = CTR; c.border = BDR
        ws.column_dimensions[get_column_letter(ci)].width = width
    ws.row_dimensions[2].height = 20

    # Datos
    for ri, row in enumerate(rows, 3):
        for ci, (key, _, _) in enumerate(use_cols, 1):
            val = row.get(key, "—")
            c   = ws.cell(row=ri, column=ci, value=val)
            c.font = D_FONT; c.border = BDR
            c.alignment = CTR if key in ("id","year","value") else LEFT
            if key == "value" and isinstance(val, float):
                c.number_format = "#,##0"
            if ri % 2 == 0: c.fill = ALT
        ws.row_dimensions[ri].height = 16

    return wb


def _generate_excel_multisheet(rows, title):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()

    # ── Estilos ──────────────────────────────────────────────────────
    H_FILL  = PatternFill("solid", start_color="1E3A6E")
    H_FONT  = Font(bold=True, color="FFFFFF", name="Arial", size=10)
    T_FILL  = PatternFill("solid", start_color="B91C3C")
    T_FONT  = Font(bold=True, color="FFFFFF", name="Arial", size=12)
    D_FONT  = Font(name="Arial", size=9)
    ALT     = PatternFill("solid", start_color="EFF6FF")
    thin    = Side(style="thin", color="D1D5DB")
    BORDER  = Border(left=thin, right=thin, top=thin, bottom=thin)
    CENTER  = Alignment(horizontal="center", vertical="center", wrap_text=True)
    LEFT    = Alignment(horizontal="left",   vertical="center", wrap_text=True)
    MONEY   = "#,##0"
    DATE_FMT= "DD/MM/YYYY"

    def title_row(ws, text, n_cols, color="B91C3C"):
        ws.merge_cells(f"A1:{get_column_letter(n_cols)}1")
        c = ws["A1"]
        c.value     = text
        c.font      = Font(bold=True, size=12, color="FFFFFF", name="Arial")
        c.fill      = PatternFill("solid", start_color=color)
        c.alignment = CENTER
        ws.row_dimensions[1].height = 24

    def header_row(ws, cols, row=2):
        for ci, (_, label, width) in enumerate(cols, 1):
            c = ws.cell(row=row, column=ci, value=label)
            c.font = H_FONT; c.fill = H_FILL
            c.alignment = CENTER; c.border = BORDER
            ws.column_dimensions[get_column_letter(ci)].width = width
        ws.row_dimensions[row].height = 22

    def data_row(ws, ri, values, money_cols=(), date_cols=(), center_cols=()):
        for ci, val in enumerate(values, 1):
            c = ws.cell(row=ri, column=ci, value=val)
            c.font   = D_FONT
            c.border = BORDER
            c.alignment = CENTER if ci in center_cols else LEFT
            if ci in money_cols and isinstance(val, (int, float)):
                c.number_format = MONEY
            if ci in date_cols and val and val not in ("N/A","—",""):
                c.number_format = DATE_FMT
            if ri % 2 == 0:
                c.fill = ALT
        ws.row_dimensions[ri].height = 16

    # ════════════════════════════════════════════════════════════════
    # HOJA 1 — RESUMEN GENERAL
    # ════════════════════════════════════════════════════════════════
    ws1 = wb.active
    ws1.title = "Resumen General"
    ws1.freeze_panes = "A3"

    COLS1 = [
        ("id",           "#ID",                   7),
        ("year",         "Año",                   6),
        ("external",     "N° Externo",            16),
        ("name",         "Nombre del Proyecto",   48),
        ("purpose",      "Objeto",                48),
        ("status",       "Estado",                18),
        ("entity",       "Entidad",               32),
        ("department",   "Dependencia",           26),
        ("value",        "Valor Contrato",        18),
        ("entity_contribution", "Aporte Entidad", 18),
        ("univ_contribution",   "Aporte UD",      16),
        ("benefit_pct",  "% Beneficio",           12),
        ("benefit",      "Beneficio Inst.",        20),
        ("start_date",   "Fecha Inicio",          14),
        ("end_date",     "Fecha Fin",             14),
        ("sub_date",     "Fecha Suscripción",      16),
        ("supervisor",   "Supervisor",             16),
        ("email",        "Correo Principal",       28),
        ("act",          "Acto Administrativo",    20),
    ]

    title_row(ws1, title, len(COLS1), "1E3A6E")
    header_row(ws1, COLS1)

    for i, row in enumerate(rows, 3):
        vals = [
            row.get("id"),     row.get("year"),   row.get("external"),
            row.get("name"),   row.get("purpose"),
            row.get("status"), row.get("entity"), row.get("department"),
            float(row.get("value", 0)),
            float(row.get("entity_contribution", 0)),
            float(row.get("university_contribution", 0)),
            float(row.get("benefit_pct", 12)),
            float(row.get("benefit", 0)),
            row.get("start_date"), row.get("end_date"), row.get("sub_date"),
            row.get("supervisor"), row.get("email"), row.get("act"),
        ]
        data_row(ws1, i, vals,
                 money_cols=(9,10,11,13), date_cols=(14,15,16),
                 center_cols=(1,2,12))

    # Totales
    tr = len(rows) + 3
    ws1.cell(row=tr, column=1, value="TOTAL").font = Font(bold=True, name="Arial", size=10)
    for col, letter in [(9,"I"),(10,"J"),(11,"K"),(13,"M")]:
        c = ws1.cell(row=tr, column=col, value=f"=SUM({letter}3:{letter}{tr-1})")
        c.font = Font(bold=True, name="Arial", size=10)
        c.number_format = MONEY; c.border = BORDER; c.alignment = CENTER

    # ════════════════════════════════════════════════════════════════
    # HOJA 2 — MODIFICACIONES  (se llena desde el endpoint con datos reales)
    # ════════════════════════════════════════════════════════════════
    ws2 = wb.create_sheet("Modificaciones")
    ws2.freeze_panes = "A3"

    COLS2 = [
        ("n",       "#",                    5),
        ("pid",     "Project ID",           10),
        ("year",    "Año",                   6),
        ("ext",     "N° Externo",           16),
        ("mod_num", "N° Mod.",               8),
        ("tipo",    "Tipo",                 22),
        ("fecha",   "Fecha Aprobación",     16),
        ("acto",    "Acto Administrativo",  20),
        ("adicion", "Valor Adición",        18),
        ("nueva_fin","Nueva Fecha Fin",     16),
        ("dias",    "Días Prórroga",         12),
        ("nuevo_val","Nuevo Valor Total",   18),
        ("estado",  "Estado",               10),
        ("justif",  "Justificación",        40),
    ]
    title_row(ws2, "MODIFICACIONES DE PROYECTOS", len(COLS2), "0F2952")
    header_row(ws2, COLS2)

    TIPO_LABEL = {
        "ADDITION":"Adición","EXTENSION":"Prórroga","BOTH":"Adición + Prórroga",
        "CONTRACTUAL":"Modificación Contractual","SUSPENSION":"Suspensión",
        "RESTART":"Reinicio","CESION_CESIONARIA":"Cesión (Cesionaria)",
        "CESION_CEDENTE":"Cesión (Cedente)","LIQUIDATION":"Liquidación",
    }
    proj_map = {r["id"]: r for r in rows}

    # Las modificaciones se inyectan desde download_excel (ver abajo)
    ws2._siexud_proj_map = proj_map
    ws2._siexud_COLS2    = COLS2
    ws2._siexud_TIPO     = TIPO_LABEL

    # ════════════════════════════════════════════════════════════════
    # HOJA 3 — DOCUMENTOS
    # ════════════════════════════════════════════════════════════════
    ws3 = wb.create_sheet("Documentos")
    ws3.freeze_panes = "A3"

    COLS3 = [
        ("n",       "#",                5),
        ("pid",     "Project ID",      10),
        ("year",    "Año",              6),
        ("ext",     "N° Externo",      16),
        ("doc_num", "N° Doc.",          8),
        ("type",    "Tipo",            18),
        ("fname",   "Nombre Archivo",  40),
        ("fecha",   "Fecha",           14),
        ("size",    "Tamaño (KB)",     12),
        ("obs",     "Observaciones",   40),
    ]
    title_row(ws3, "DOCUMENTOS CARGADOS POR PROYECTO", len(COLS3), "6366F1")
    header_row(ws3, COLS3)
    ws3._siexud_proj_map = proj_map
    ws3._siexud_COLS3    = COLS3

    # ════════════════════════════════════════════════════════════════
    # HOJA 4 — RESUMEN FINANCIERO
    # ════════════════════════════════════════════════════════════════
    ws4 = wb.create_sheet("Resumen Financiero")

    title_row(ws4, "RESUMEN FINANCIERO — SIEXUD", 3, "10B981")

    ws4.column_dimensions["A"].width = 36
    ws4.column_dimensions["B"].width = 24
    ws4.column_dimensions["C"].width = 14

    meta_items = [
        ("Fecha de generación:", date.today().strftime("%d/%m/%Y"), ""),
        ("Total proyectos en el reporte:", len(rows), ""),
        ("", "", ""),
    ]
    for ri, (k, v, _) in enumerate(meta_items, 3):
        ws4.cell(row=ri, column=1, value=k).font  = Font(bold=True, name="Arial", size=10)
        ws4.cell(row=ri, column=2, value=v).font   = Font(name="Arial", size=10)

    # Tabla financiera
    start_row = 7
    for ci, label in enumerate(["Métrica", "Valor Total (COP)", "# Proyectos"], 1):
        c = ws4.cell(row=start_row, column=ci, value=label)
        c.font = H_FONT; c.fill = H_FILL; c.alignment = CENTER; c.border = BORDER
    ws4.row_dimensions[start_row].height = 22

    # Calcular métricas
    total_val  = sum(float(r.get("value", 0)) for r in rows)
    total_ent  = sum(float(r.get("entity_contribution", 0)) for r in rows)
    total_univ = sum(float(r.get("university_contribution", 0)) for r in rows)
    total_ben  = sum(float(r.get("benefit", 0)) for r in rows)

    by_status = {}
    for r in rows:
        st = r.get("status", "Sin estado")
        by_status.setdefault(st, {"value": 0, "count": 0})
        by_status[st]["value"] += float(r.get("value", 0))
        by_status[st]["count"] += 1

    fin_rows = [
        ("Valor total de contratos",      total_val,  len(rows)),
        ("Total aporte entidades",         total_ent,  len(rows)),
        ("Total aporte Universidad",       total_univ, len(rows)),
        ("Total beneficio institucional",  total_ben,  len(rows)),
        ("", "", ""),
    ]
    for st_name, st_data in sorted(by_status.items()):
        fin_rows.append((f"  Estado: {st_name}", st_data["value"], st_data["count"]))

    for i, (label, val, cnt) in enumerate(fin_rows):
        ri = start_row + 1 + i
        c1 = ws4.cell(row=ri, column=1, value=label)
        c2 = ws4.cell(row=ri, column=2, value=val if val != "" else "")
        c3 = ws4.cell(row=ri, column=3, value=cnt if cnt != "" else "")
        for c in (c1, c2, c3):
            c.font = D_FONT; c.border = BORDER if label else None; c.alignment = LEFT
        if isinstance(val, float):
            c2.number_format = MONEY
            c2.alignment = CENTER
        if isinstance(cnt, int):
            c3.alignment = CENTER
        if i % 2 == 0 and label:
            c1.fill = ALT; c2.fill = ALT; c3.fill = ALT

    return wb   # retorna el workbook, no bytes — para que download_excel pueda agregar mods y docs


def generate_pdf(rows, title, columns=None):
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    )

    UD_RED  = colors.HexColor('#B91C3C')
    UD_BLUE = colors.HexColor('#1E3A6E')
    LIGHT   = colors.HexColor('#EFF6FF')
    GRAY    = colors.HexColor('#64748B')

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(A4),
                            leftMargin=1.5*cm, rightMargin=1.5*cm,
                            topMargin=1.5*cm, bottomMargin=1.5*cm)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('title', fontSize=16, fontName='Helvetica-Bold',
                                  textColor=UD_BLUE, spaceAfter=4)
    sub_style   = ParagraphStyle('sub',   fontSize=9,  fontName='Helvetica',
                                  textColor=GRAY, spaceAfter=12)
    cell_style  = ParagraphStyle('cell',  fontSize=7.5, fontName='Helvetica',
                                  leading=10, wordWrap='LTR')
    head_style  = ParagraphStyle('head',  fontSize=8, fontName='Helvetica-Bold',
                                  textColor=colors.white, alignment=1)

    story = []

    # Header
    story.append(Paragraph("SIEXUD — Universidad Distrital Francisco José de Caldas", sub_style))
    story.append(Paragraph(title, title_style))
    story.append(Paragraph(
        f"Generado el {date.today().strftime('%d/%m/%Y')} · Total: {len(rows)} proyectos",
        sub_style
    ))
    story.append(HRFlowable(width="100%", thickness=2, color=UD_RED, spaceAfter=10))

    # Definir columnas
    ALL_COLS = [
        ('year',       'Año',       1.2*cm),
        ('external',   'N° Ext.',   2.5*cm),
        ('name',       'Nombre',    7*cm),
        ('purpose',    'Objeto',    7*cm),
        ('status',     'Estado',    3*cm),
        ('entity',     'Entidad',   5*cm),
        ('value',      'Valor',     3*cm),
        ('start_date', 'Inicio',    2.2*cm),
        ('end_date',   'Fin',       2.2*cm),
    ]

    if columns:
        col_map = {
            'year': 'year', 'number': 'external', 'purpose': 'purpose',
            'name': 'name', 'entity': 'entity', 'value': 'value',
            'status': 'status', 'department': 'department',
        }
        keys = set()
        for c in columns:
            k = col_map.get(c)
            if k: keys.add(k)
        if keys:
            ALL_COLS = [c for c in ALL_COLS if c[0] in keys]

    # Tabla
    col_widths = [c[2] for c in ALL_COLS]
    header_row = [Paragraph(c[1], head_style) for c in ALL_COLS]
    table_data = [header_row]

    for row in rows:
        data_row = []
        for key, _, _ in ALL_COLS:
            val = row.get(key, '—')
            if key == 'value' and isinstance(val, float):
                val = f"${val:,.0f}".replace(',', '.')
            text = str(val)[:120] if val else '—'
            data_row.append(Paragraph(text, cell_style))
        table_data.append(data_row)

    tbl = Table(table_data, colWidths=col_widths, repeatRows=1)
    tbl.setStyle(TableStyle([
        ('BACKGROUND',  (0,0), (-1,0), UD_BLUE),
        ('TEXTCOLOR',   (0,0), (-1,0), colors.white),
        ('FONTNAME',    (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE',    (0,0), (-1,0), 8),
        ('ALIGN',       (0,0), (-1,-1), 'LEFT'),
        ('VALIGN',      (0,0), (-1,-1), 'TOP'),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, LIGHT]),
        ('GRID',        (0,0), (-1,-1), 0.3, colors.HexColor('#CBD5E1')),
        ('TOPPADDING',  (0,0), (-1,-1), 4),
        ('BOTTOMPADDING',(0,0), (-1,-1), 4),
        ('LEFTPADDING', (0,0), (-1,-1), 4),
    ]))
    story.append(tbl)

    doc.build(story)
    buf.seek(0)
    return buf.read()


def generate_word(rows, title, columns=None):
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        raise HTTPException(500, "python-docx no instalado. Agregar al requirements.txt")

    ALL_COLS = [
        ('year',       'Año'),
        ('external',   'N° Externo'),
        ('name',       'Nombre del Proyecto'),
        ('purpose',    'Objeto'),
        ('status',     'Estado'),
        ('entity',     'Entidad'),
        ('value',      'Valor Contrato'),
        ('start_date', 'Fecha Inicio'),
        ('end_date',   'Fecha Fin'),
    ]

    if columns:
        col_map = {
            'year': 'year', 'number': 'external', 'purpose': 'purpose',
            'name': 'name', 'entity': 'entity', 'value': 'value',
            'status': 'status',
        }
        keys = set()
        for c in columns:
            k = col_map.get(c)
            if k: keys.add(k)
        if keys:
            ALL_COLS = [c for c in ALL_COLS if c[0] in keys]

    doc = DocxDocument()

    # Título
    h = doc.add_heading(title, 0)
    h.runs[0].font.color.rgb = RGBColor(30, 58, 110)

    doc.add_paragraph(
        f"Universidad Distrital Francisco José de Caldas · SIEXUD\n"
        f"Generado: {date.today().strftime('%d/%m/%Y')} · Total proyectos: {len(rows)}"
    ).runs[0].font.size = Pt(9)

    # Tabla
    table = doc.add_table(rows=1, cols=len(ALL_COLS))
    table.style = 'Table Grid'

    # Encabezado
    hdr_cells = table.rows[0].cells
    for i, (_, label) in enumerate(ALL_COLS):
        cell = hdr_cells[i]
        cell.text = label
        run = cell.paragraphs[0].runs[0]
        run.font.bold  = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size  = Pt(9)
        cell._tc.get_or_add_tcPr()

    # Datos
    for row in rows:
        cells = table.add_row().cells
        for i, (key, _) in enumerate(ALL_COLS):
            val = row.get(key, '—')
            if key == 'value' and isinstance(val, float):
                val = f"${val:,.0f}".replace(',', '.')
            cells[i].text = str(val)[:200] if val else '—'
            cells[i].paragraphs[0].runs[0].font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── GET /reportes/stats/ ─────────────────────────────────────────────
@router.get("/stats/")
def get_stats(
    year:                Optional[str] = Query(None),
    status_ids:          Optional[str] = Query(None),
    valor_min:           Optional[str] = Query(None),
    valor_max:           Optional[str] = Query(None),
    fecha_inicio_desde:  Optional[str] = Query(None),
    fecha_inicio_hasta:  Optional[str] = Query(None),
    fecha_fin_desde:     Optional[str] = Query(None),
    fecha_fin_hasta:     Optional[str] = Query(None),
    entidad_id:          Optional[str] = Query(None),
    departamento_id:     Optional[str] = Query(None),
    db: Session = Depends(get_db),
):
    q = db.query(Project).filter(Project.is_active == True)
    q = apply_filters(q, year, status_ids, valor_min, valor_max,
                      fecha_inicio_desde, fecha_inicio_hasta,
                      fecha_fin_desde, fecha_fin_hasta,
                      entidad_id, departamento_id)
    proyectos = q.all()
    ids = [p.project_id for p in proyectos]

    con_mods = 0
    if ids:
        con_mods = db.query(func.count(func.distinct(ProjectModification.project_id))) \
            .filter(ProjectModification.project_id.in_(ids), ProjectModification.is_active == True).scalar() or 0

    por_estado = {}
    for p in proyectos:
        nombre = get_status_name(db, p.project_status_id)
        por_estado[nombre] = por_estado.get(nombre, 0) + 1

    por_anio = {}
    for p in proyectos:
        anio = str(p.project_year)
        por_anio[anio] = por_anio.get(anio, 0) + 1

    return {
        "total_proyectos":    len(proyectos),
        "valor_total":        sum(float(p.project_value or 0) for p in proyectos),
        "beneficio_total":    sum(float(p.institutional_benefit_value or 0) for p in proyectos),
        "con_modificaciones": con_mods,
        "por_estado":         por_estado,
        "por_anio":           por_anio,
    }


# ── GET /reportes/excel/ ─────────────────────────────────────────────
@router.get("/excel/")
def download_excel(
    year:                Optional[str] = Query(None),
    status_ids:          Optional[str] = Query(None),
    valor_min:           Optional[str] = Query(None),
    valor_max:           Optional[str] = Query(None),
    fecha_inicio_desde:  Optional[str] = Query(None),
    fecha_inicio_hasta:  Optional[str] = Query(None),
    fecha_fin_desde:     Optional[str] = Query(None),
    fecha_fin_hasta:     Optional[str] = Query(None),
    entidad_id:          Optional[str] = Query(None),
    departamento_id:     Optional[str] = Query(None),
    db: Session = Depends(get_db),
):
    q = db.query(Project).filter(Project.is_active == True)
    q = apply_filters(q, year, status_ids, valor_min, valor_max,
                      fecha_inicio_desde, fecha_inicio_hasta,
                      fecha_fin_desde, fecha_fin_hasta,
                      entidad_id, departamento_id)
    proyectos = q.order_by(Project.project_year.desc(), Project.project_id).all()

    if not proyectos:
        raise HTTPException(404, "No hay proyectos con los filtros seleccionados.")

    rows = build_project_rows(db, proyectos)

    # Títulos con filtros
    title_parts = ["REPORTE GLOBAL DE PROYECTOS SIEXUD"]
    if year:       title_parts.append(f"AÑO {year}")
    if status_ids: title_parts.append(f"ESTADOS FILTRADOS")
    title = " — ".join(title_parts) + f" · {date.today().strftime('%d/%m/%Y')}"

    # Generar workbook con 4 hojas
    wb = generate_excel(rows, title)
    project_ids = [p.project_id for p in proyectos]

    # ── Llenar Hoja 2: Modificaciones ──────────────────────────────
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    thin   = Side(style="thin", color="D1D5DB")
    BORDER = Border(left=thin, right=thin, top=thin, bottom=thin)
    D_FONT = Font(name="Arial", size=9)
    ALT    = PatternFill("solid", start_color="EFF6FF")
    LEFT   = Alignment(horizontal="left", vertical="center", wrap_text=True)
    CENTER = Alignment(horizontal="center", vertical="center")
    MONEY  = "#,##0"
    DATE_FMT = "DD/MM/YYYY"

    ws2 = wb["Modificaciones"]
    proj_map = {p.project_id: p for p in proyectos}

    TIPO_LABEL = {
        "ADDITION":"Adición","EXTENSION":"Prórroga","BOTH":"Adición + Prórroga",
        "CONTRACTUAL":"Modificación Contractual","SUSPENSION":"Suspensión",
        "RESTART":"Reinicio","CESION_CESIONARIA":"Cesión (Cesionaria)",
        "CESION_CEDENTE":"Cesión (Cedente)","LIQUIDATION":"Liquidación",
    }

    mods = db.query(ProjectModification).filter(
        ProjectModification.project_id.in_(project_ids),
    ).order_by(ProjectModification.project_id, ProjectModification.modification_number).all()

    for i, m in enumerate(mods, 3):
        p = proj_map.get(m.project_id)
        vals = [
            i-2,
            m.project_id,
            p.project_year if p else "—",
            p.external_project_number if p else "—",
            m.modification_number,
            TIPO_LABEL.get(m.modification_type, m.modification_type),
            m.approval_date,
            m.administrative_act or "—",
            float(m.addition_value or 0) if m.addition_value else "—",
            m.new_end_date,
            m.extension_days or "—",
            float(m.new_total_value or 0) if m.new_total_value else "—",
            "Activo" if m.is_active else "Inactivo",
            (m.justification or "—")[:200],
        ]
        for ci, val in enumerate(vals, 1):
            c = ws2.cell(row=i, column=ci, value=val)
            c.font = D_FONT; c.border = BORDER; c.alignment = LEFT
            if ci in (9, 12) and isinstance(val, float):
                c.number_format = MONEY; c.alignment = CENTER
            if ci in (7, 10) and val and val not in ("—", ""):
                c.number_format = DATE_FMT
            if i % 2 == 0: c.fill = ALT
        ws2.row_dimensions[i].height = 16

    if not mods:
        ws2.cell(row=3, column=1, value="No hay modificaciones para los proyectos seleccionados").font = Font(italic=True, name="Arial", size=10, color="94A3B8")

    # ── Llenar Hoja 3: Documentos ───────────────────────────────────
    ws3 = wb["Documentos"]
    docs = db.query(ProjectDocument).filter(
        ProjectDocument.project_id.in_(project_ids),
        ProjectDocument.is_active == True,
    ).order_by(ProjectDocument.project_id, ProjectDocument.document_number).all()

    for i, d in enumerate(docs, 3):
        p = proj_map.get(d.project_id)
        vals = [
            i-2,
            d.project_id,
            p.project_year if p else "—",
            p.external_project_number if p else "—",
            d.document_number,
            d.type_code if hasattr(d, "type_code") else str(d.document_type_id),
            d.original_filename or d.document_name,
            d.document_date,
            round(d.file_size / 1024, 1) if d.file_size else "—",
            (d.observations or "—")[:150],
        ]
        for ci, val in enumerate(vals, 1):
            c = ws3.cell(row=i, column=ci, value=val)
            c.font = D_FONT; c.border = BORDER; c.alignment = LEFT
            if ci == 8 and val and val not in ("—", ""):
                c.number_format = DATE_FMT
            if i % 2 == 0: c.fill = ALT
        ws3.row_dimensions[i].height = 16

    if not docs:
        ws3.cell(row=3, column=1, value="No hay documentos cargados para los proyectos seleccionados").font = Font(italic=True, name="Arial", size=10, color="94A3B8")

    # ── Serializar ──────────────────────────────────────────────────
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    file_bytes = buf.read()

    filename = f"SIEXUD_Reporte_{date.today().isoformat()}.xlsx"
    return Response(
        content=file_bytes,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── GET /reportes/download/{token} ───────────────────────────────────
@router.get("/download/{token}")
def download_file(token: str):
    """Descarga un archivo generado por el chat IA."""
    if token not in _file_store:
        raise HTTPException(404, "Archivo no encontrado o expirado. Vuelve a generar el reporte.")
    entry = _file_store[token]
    # Verificar expiración
    now = datetime.utcnow().timestamp()
    if now - entry.get("created_at", 0) > _FILE_TTL_SECONDS:
        del _file_store[token]
        raise HTTPException(410, "El archivo expiró (30 min). Vuelve a hacer la consulta para regenerarlo.")
    return Response(
        content=entry['content'],
        media_type=entry['mimetype'],
        headers={"Content-Disposition": f'attachment; filename="{entry["filename"]}"'},
    )


# ── POST /reportes/chat/ ─────────────────────────────────────────────
class ChatRequest(BaseModel):
    message:  str
    context:  dict = {}
    history:  list = []  # historial previo para contexto multi-turno


@router.post("/chat/")
async def chat_ia(data: ChatRequest, db: Session = Depends(get_db)):
    api_key = getattr(settings, 'anthropic_api_key', None)
    if not api_key:
        raise HTTPException(500, "ANTHROPIC_API_KEY no configurada en .env")

    # ── 1. Detectar intención ─────────────────────────────────────
    intent = detect_intent(data.message)

    download_url = None
    file_context = ""

    # ── 2. Si pide un archivo, generarlo ─────────────────────────
    if intent['is_report'] and intent['file_type']:
        try:
            proyectos = get_projects_for_report(db, intent['filters'], data.context.get('filters'))
            rows = build_project_rows(db, proyectos, intent.get('columns'))

            ftype = intent['file_type']
            year_label  = intent['filters'].get('year', '')
            status_label = intent['filters'].get('status_name', '')
            title_parts = ['REPORTE DE PROYECTOS SIEXUD']
            if year_label:   title_parts.append(f"AÑO {year_label}")
            if status_label: title_parts.append(status_label.upper())
            title = ' — '.join(title_parts)

            if ftype == 'excel':
                wb = generate_excel(rows, title, intent.get('columns'))
                # generate_excel retorna Workbook — serializar a bytes
                _buf = io.BytesIO()
                wb.save(_buf)
                _buf.seek(0)
                content  = _buf.read()
                mimetype = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                ext      = "xlsx"
            elif ftype == 'pdf':
                content  = generate_pdf(rows, title, intent.get('columns'))
                mimetype = "application/pdf"
                ext      = "pdf"
            elif ftype == 'word':
                content  = generate_word(rows, title, intent.get('columns'))
                mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                ext      = "docx"
            else:
                content  = generate_pdf(rows, title, intent.get('columns'))
                mimetype = "application/pdf"
                ext      = "pdf"

            # Guardar en store temporal
            token    = str(uuid.uuid4()).replace('-', '')[:16]
            filename = f"SIEXUD_{year_label or 'todos'}_{date.today().isoformat()}.{ext}"
            _file_store[token] = {
                "content":    content,
                "mimetype":   mimetype,
                "filename":   filename,
                "created_at": datetime.utcnow().timestamp(),
            }
            # Limpiar archivos expirados (> 30 min)
            now = datetime.utcnow().timestamp()
            expired = [k for k, v in _file_store.items() if now - v.get("created_at", 0) > _FILE_TTL_SECONDS]
            for k in expired:
                del _file_store[k]

            download_url = f"/reportes/download/{token}"

            # Contexto para Claude
            file_context = f"""
ARCHIVO GENERADO EXITOSAMENTE:
- Tipo: {ext.upper()}
- Nombre: {filename}
- Proyectos incluidos: {len(rows)}
- Filtros aplicados: {json.dumps(intent['filters'], ensure_ascii=False)}
- URL de descarga: {download_url}

Dile al usuario que el archivo está listo e incluye EXACTAMENTE esta línea al final de tu respuesta (sin modificar):
[DESCARGA]:{download_url}
"""
        except Exception as e:
            file_context = f"ERROR al generar archivo: {str(e)}. Informa al usuario del error."

    # ── 3. Contexto de BD (optimizado por tokens) ─────────────────
    # Solo cargamos datos detallados si son necesarios
    msg_lower = data.message.lower()
    need_detail = any(w in msg_lower for w in [
        'cuántos', 'cuantos', 'total', 'suma', 'promedio', 'lista',
        'entidad', 'estado', 'año', 'valor', 'qué proyectos', 'que proyectos'
    ])

    # Siempre: totales generales (muy pocos tokens)
    total_proy  = db.query(func.count(Project.project_id)).filter(Project.is_active == True).scalar()
    total_valor = db.query(func.sum(Project.project_value)).filter(Project.is_active == True).scalar()
    total_mods  = db.query(func.count(ProjectModification.modification_id)).filter(ProjectModification.is_active == True).scalar()
    total_docs  = db.query(func.count(ProjectDocument.document_id)).filter(ProjectDocument.is_active == True).scalar()

    # Condicional: desglose por año y estado
    detail_context = ""
    if need_detail:
        por_anio = db.query(Project.project_year, func.count(Project.project_id)) \
            .filter(Project.is_active == True) \
            .group_by(Project.project_year) \
            .order_by(Project.project_year.desc()).limit(15).all()

        por_estado = db.query(ProjectStatus.status_name, func.count(Project.project_id)) \
            .join(Project, Project.project_status_id == ProjectStatus.status_id) \
            .filter(Project.is_active == True) \
            .group_by(ProjectStatus.status_name).all()

        por_entidad = db.query(Entity.entity_name, func.count(Project.project_id)) \
            .join(Project, Project.entity_id == Entity.entity_id) \
            .filter(Project.is_active == True) \
            .group_by(Entity.entity_name) \
            .order_by(func.count(Project.project_id).desc()).limit(8).all()

        detail_context = f"""
PROYECTOS POR AÑO: {', '.join(f'{a}:{c}' for a, c in por_anio)}
PROYECTOS POR ESTADO: {', '.join(f'{e}:{c}' for e, c in por_estado)}
TOP ENTIDADES: {', '.join(f'{e}:{c}' for e, c in por_entidad)}"""

    db_context = f"""SIEXUD - Base de datos actual:
- Total proyectos activos: {total_proy}
- Valor total contratos: {fmt_money(total_valor)} COP
- Modificaciones registradas: {total_mods}
- Documentos cargados: {total_docs}{detail_context}
{file_context}"""

    # ── 4. System prompt optimizado ──────────────────────────────
    system_prompt = """Eres el asistente IA del sistema SIEXUD (Universidad Distrital Francisco José de Caldas).
Responde SIEMPRE en español, de forma concisa y útil.
Puedes generar reportes en Excel, PDF o Word cuando el usuario los solicite.
Si el contexto incluye una URL de descarga con el formato [DESCARGA]:/reportes/download/XXX, inclúyela en tu respuesta exactamente así para que el frontend la detecte y muestre el botón de descarga.
Usa solo los datos del contexto. No inventes información.
Si no tienes el dato exacto, dilo claramente."""

    # ── 5. Historial multi-turno (máx 6 turnos para ahorrar tokens) ──
    messages = []
    for h in data.history[-6:]:
        if h.get('role') in ('user', 'assistant') and h.get('content'):
            messages.append({"role": h['role'], "content": str(h['content'])[:500]})

    messages.append({"role": "user", "content": f"{db_context}\n\nPregunta: {data.message}"})

    # ── 6. Llamar a Claude ────────────────────────────────────────
    try:
        async with httpx.AsyncClient(timeout=40.0, verify=False) as client:
            response = await client.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "x-api-key":          api_key,
                    "anthropic-version":  "2023-06-01",
                    "content-type":       "application/json",
                },
                json={
                    "model":      "claude-haiku-4-5-20251001",
                    "max_tokens": 800,
                    "system":     system_prompt,
                    "messages":   messages,
                },
            )
        if response.status_code != 200:
            raise HTTPException(500, f"Error API Claude: {response.text}")

        answer = response.json()["content"][0]["text"]
        return {
            "response":     answer,
            "download_url": download_url,
            "file_type":    intent['file_type'],
            "projects_count": len(rows) if intent['is_report'] and 'rows' in dir() else None,
        }

    except httpx.TimeoutException:
        raise HTTPException(504, "Timeout al conectar con Claude API")
    except Exception as e:
        raise HTTPException(500, f"Error IA: {str(e)}")
